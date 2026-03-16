from flask import Flask, render_template, jsonify, send_from_directory, request
import os
import re
import yaml
import socket
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['STATIC_FOLDER'] = 'static'
#app.config['LOG_FILE_PATH'] = 'static/engine.log'
app.config['LOG_FILE_PATH'] = os.path.expanduser('~/.config/Z-Factory/logs/') + datetime.now().strftime('%Y-%m-%d') + '/engine_logs/engine.log'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['CONFIG_FILE'] = 'config.yaml'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['STATIC_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def load_config():
    config = {
        'mail': {
            'server': 'smtp.gmail.com',
            'port': 587,
            'use_tls': True,
            'from_email': '',
            'password': ''
        },
        'recipients': [],
        'report': {
            'title': 'Engine Logs 分析报告',
            'company': '您的公司名称'
        }
    }

    if os.path.exists(app.config['CONFIG_FILE']):
        try:
            with open(app.config['CONFIG_FILE'], 'r', encoding='utf-8') as f:
                yaml_config = yaml.safe_load(f)
                if yaml_config:
                    if 'mail' in yaml_config:
                        config['mail'].update(yaml_config['mail'])
                    if 'recipients' in yaml_config:
                        config['recipients'] = yaml_config['recipients']
                    if 'report' in yaml_config:
                        config['report'].update(yaml_config['report'])
        except Exception as e:
            print(f'加载配置文件失败: {e}')

    return config

app.config.update(load_config())

def get_local_ip():
    """获取本机局域网IP地址"""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return '127.0.0.1'

class LogAnalyzer:
    def __init__(self, log_content):
        self.log_content = log_content
        self.test_cases = []
        self.log_lines = log_content.split('\n')

    def analyze(self):
        current_test_case = None
        current_lines = []
        in_test_case = False

        for idx, line in enumerate(self.log_lines):
            match_start = re.search(r'log_service\.py.*?Log\.Info\s*:\s*测试用例(\d+)\s*开始', line)
            match_end = re.search(r'log_service\.py.*?Log\.Info\s*:\s*测试用例(\d+)\s*结束', line)

            if match_start:
                test_case_id = match_start.group(1)

                if in_test_case and current_test_case:
                    self._process_test_case(current_test_case, current_lines)

                current_test_case = {
                    'test_case_id': test_case_id,
                    'component_info': {},
                    'has_error': False,
                    'errors': [],
                    'start_line': idx
                }
                current_lines = [(idx, line)]
                in_test_case = True

            elif match_end and in_test_case and current_test_case:
                end_test_case_id = match_end.group(1)
                if end_test_case_id == current_test_case['test_case_id']:
                    current_lines.append((idx, line))
                    self._process_test_case(current_test_case, current_lines)
                    current_test_case = None
                    in_test_case = False

            elif in_test_case:
                current_lines.append((idx, line))

        if in_test_case and current_test_case:
            self._process_test_case(current_test_case, current_lines)

        return self.test_cases

    def _process_test_case(self, test_case, lines):
        component_info = {
            'module': '',
            'name_cn': '',
            'category': '',
            'method': '',
            'plugin_name': '',
            'unique_id': ''
        }

        for line_idx, line in lines:
            if 'ERROR' in line:
                # 从ERROR或Exception开始提取错误信息
                error_match = re.search(r'(ERROR|Exception|Error)(.*)', line)
                if error_match:
                    error_message = (error_match.group(1) + error_match.group(2)).strip()
                    test_case['has_error'] = True
                    test_case['errors'].append({
                        'line': line_idx,
                        'message': error_message
                    })

            for key, pattern in {
                'module': r'组件模块为([^|]+)',
                'name_cn': r'组件中文名为([^|]+)',
                'category': r'组件分类为([^|]+)',
                'method': r'组件方法名为([^|]+)',
                'plugin_name': r'组件插件包中文名为([^|]+)',
                'unique_id': r'组件唯一标识为([^|]+)'
            }.items():
                if not component_info.get(key):
                    match = re.search(pattern, line)
                    if match and match.group(1):
                        value = match.group(1).strip()
                        if value and value != '为':
                            component_info[key] = value

        test_case['component_info'] = component_info
        self.test_cases.append(test_case)

def generate_word_report(test_cases, output_path):
    doc = Document()

    title = doc.add_heading(app.config['report']['title'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_time = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
    doc.add_paragraph(f'报告生成时间: {report_time}')
    doc.add_paragraph(f'生成单位: {app.config["report"]["company"]}')

    doc.add_paragraph()

    total = len(test_cases)
    errors = sum(1 for tc in test_cases if tc['has_error'])
    normal = total - errors

    summary = doc.add_heading('一、总体统计', level=1)

    summary_info = doc.add_paragraph()
    summary_info.add_run('测试用例总数: ').bold = True
    summary_info.add_run(str(total))

    normal_para = doc.add_paragraph()
    normal_para.add_run('正常用例: ').bold = True
    normal_run = normal_para.add_run(str(normal))
    normal_run.font.color.rgb = RGBColor(0x00, 0xC0, 0x00)

    error_para = doc.add_paragraph()
    error_para.add_run('异常用例: ').bold = True
    error_run = error_para.add_run(str(errors))
    error_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    detail = doc.add_heading('二、测试用例详情', level=1)

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = '测试用例ID'
    header_cells[1].text = '组件模块'
    header_cells[2].text = '组件中文名'
    header_cells[3].text = '组件分类'
    header_cells[4].text = '组件方法名'
    header_cells[5].text = '组件插件包中文名'
    header_cells[6].text = '组件唯一标识'
    header_cells[7].text = '运行结果'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for test_case in test_cases:
        row_cells = table.add_row().cells
        comp = test_case['component_info']

        row_cells[0].text = f'测试用例{test_case["test_case_id"]}'
        row_cells[0].paragraphs[0].runs[0].font.bold = True

        row_cells[1].text = comp.get('module', '')
        row_cells[2].text = comp.get('name_cn', '')
        row_cells[3].text = comp.get('category', '')
        row_cells[4].text = comp.get('method', '')
        row_cells[5].text = comp.get('plugin_name', '')
        row_cells[6].text = comp.get('unique_id', '')

        status = '正常' if not test_case['has_error'] else '异常'
        row_cells[7].text = status
        row_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if test_case['has_error']:
            row_cells[7].paragraphs[0].runs[0].font.bold = True
            row_cells[7].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            row_cells[7].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0xC0, 0x00)

    doc.save(output_path)
    return output_path

def send_email_report(recipients, word_file_path, report_url, timestamp):
    from_email = app.config['mail']['from_email']
    password = app.config['mail']['password']

    if not from_email or not password:
        return {'success': False, 'error': '邮件配置未设置，请在 config.yaml 中配置 from_email 和 password'}

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ', '.join(recipients)
    msg['Subject'] = f"{app.config['report']['title']} - {timestamp}"

    body = f'''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background: #f5f7fa;
                padding: 40px 20px;
            }}
            .container {{
                max-width: 800px;
                margin: 0 auto;
                background: white;
                border-radius: 12px;
                overflow: hidden;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px 30px;
                text-align: center;
            }}
            .header h1 {{
                font-size: 28px;
                margin-bottom: 10px;
                font-weight: 600;
            }}
            .header p {{
                opacity: 0.9;
                font-size: 14px;
            }}
            .content {{
                padding: 40px 30px;
            }}
            .greeting {{
                font-size: 16px;
                color: #333;
                margin-bottom: 20px;
                line-height: 1.6;
            }}
            .info-box {{
                background: #f8f9ff;
                border-left: 4px solid #667eea;
                padding: 20px;
                margin: 25px 0;
                border-radius: 6px;
            }}
            .info-box h3 {{
                color: #667eea;
                font-size: 18px;
                margin-bottom: 15px;
            }}
            .info-box p {{
                color: #555;
                font-size: 14px;
                line-height: 1.8;
                margin-bottom: 10px;
            }}
            .link-button {{
                display: inline-block;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 15px 30px;
                border-radius: 8px;
                text-decoration: none;
                font-weight: 600;
                margin: 20px 0;
                box-shadow: 0 4px 12px rgba(102, 126, 234, 0.3);
            }}
            .link-button:hover {{
                transform: translateY(-2px);
                box-shadow: 0 6px 16px rgba(102, 126, 234, 0.4);
            }}
            .footer {{
                background: #f8f9ff;
                padding: 20px 30px;
                text-align: center;
                border-top: 1px solid #e0e0e0;
            }}
            .footer p {{
                color: #666;
                font-size: 13px;
            }}
            .stats {{
                display: flex;
                justify-content: space-around;
                background: white;
                padding: 20px;
                border-radius: 8px;
                margin: 20px 0;
                border: 1px solid #e0e0e0;
            }}
            .stat-item {{
                text-align: center;
            }}
            .stat-value {{
                font-size: 28px;
                font-weight: bold;
                margin-bottom: 5px;
            }}
            .stat-label {{
                color: #666;
                font-size: 12px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>{app.config['report']['title']}</h1>
                <p>{app.config['report']['company']}</p>
            </div>
            
            <div class="content">
                <p class="greeting">您好，</p>
                
                <p class="greeting">Engine Logs 分析报告已生成完成，以下是本次分析的概况：</p>
                
                <div class="info-box">
                    <h3>报告概况</h3>
                    <p><strong>生成时间：</strong>{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</p>
                </div>
                
                <p style="text-align: center; color: #333; font-size: 16px;">
                    <a href="{report_url}" class="link-button">查看在线分析报告</a>
                </p>
                
                <div class="info-box">
                    <h3>附件说明</h3>
                    <p>详细的分析报告已作为附件发送给您，文件格式为 Word 文档（.docx），您可以下载后离线查看。</p>
                    <p>附件中包含：</p>
                    <ul style="margin-left: 20px; color: #555; font-size: 14px; line-height: 1.8;">
                        <li>测试用例执行概况统计</li>
                        <li>每个测试用例的详细信息</li>
                        <li>异常用例的错误信息</li>
                    </ul>
                </div>
                
                <p class="greeting">如有任何疑问，请随时联系我们。</p>
            </div>
            
            <div class="footer">
                <p>此邮件由 Engine Logs 分析工具自动发送</p>
                <p>{datetime.now().strftime('%Y年%m月%d日')}</p>
            </div>
        </div>
    </body>
    </html>
    '''

    msg.attach(MIMEText(body, 'html', 'utf-8'))

    with open(word_file_path, 'rb') as f:
        part = MIMEApplication(f.read())
        part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(word_file_path))
        msg.attach(part)

    try:
        use_ssl = app.config['mail']['port'] == 465
        if use_ssl:
            server = smtplib.SMTP_SSL(app.config['mail']['server'], app.config['mail']['port'])
        else:
            server = smtplib.SMTP(app.config['mail']['server'], app.config['mail']['port'])
            if app.config['mail']['use_tls']:
                server.starttls()
        server.login(from_email, password)
        server.send_message(msg)
        server.quit()

        return {'success': True, 'message': '邮件发送成功'}
    except Exception as e:
        return {'success': False, 'error': str(e)}

@app.route('/')
def index():
    return '<script>window.location.href="/analysis"</script>'

@app.route('/analysis')
def analysis():
    return render_template('log_analysis.html')

@app.route('/engine.log')
def get_engine_log():
    return send_from_directory(app.config['STATIC_FOLDER'], 'engine.log')

@app.route('/analyze/default')
def analyze_default_log():
    try:
        with open(app.config['LOG_FILE_PATH'], 'r', encoding='utf-8') as f:
            log_content = f.read()

        analyzer = LogAnalyzer(log_content)
        test_cases = analyzer.analyze()

        return jsonify({
            'success': True,
            'test_cases': test_cases,
            'total': len(test_cases),
            'errors': sum(1 for tc in test_cases if tc['has_error'])
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/log/context/<int:error_line>')
def get_log_context(error_line):
    try:
        with open(app.config['LOG_FILE_PATH'], 'r', encoding='utf-8') as f:
            all_lines = f.readlines()

        total_lines = len(all_lines)
        context_size = 500

        start = max(0, error_line - context_size)
        end = min(total_lines, error_line + context_size + 1)

        context_lines = []
        for i in range(start, end):
            context_lines.append({
                'line_num': i,
                'content': all_lines[i].strip(),
                'is_error': i == error_line
            })

        return jsonify({
            'lines': context_lines,
            'error_line': error_line,
            'total_lines': total_lines,
            'context_start': start,
            'context_end': end,
            'context_size': context_size
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/report/<timestamp>')
def view_report(timestamp):
    """带时间戳的报告访问路由"""
    return render_template('log_analysis.html')

@app.route('/upload', methods=['POST'])
def upload_log():
    """接收上传的engine.log文件并生成分析报告"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '未找到上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '未选择文件'}), 400

        log_content = file.read().decode('utf-8')

        analyzer = LogAnalyzer(log_content)
        test_cases = analyzer.analyze()

        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        filename = f'engine_log_report_{timestamp}.docx'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)

        generate_word_report(test_cases, output_path)

        local_ip = get_local_ip()
        report_url = f'http://{local_ip}:5000/report/{timestamp}'

        recipients = app.config.get('recipients', [])
        if recipients:
            from_email = app.config['mail']['from_email']
            password = app.config['mail']['password']

            if from_email and password:
                send_email_report(recipients, output_path, report_url, timestamp)

        return jsonify({
            'success': True,
            'report_url': report_url,
            'timestamp': timestamp,
            'total': len(test_cases),
            'errors': sum(1 for tc in test_cases if tc['has_error'])
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_startup_report():
    if not os.path.exists(app.config['LOG_FILE_PATH']):
        print(f'日志文件不存在: {app.config["LOG_FILE_PATH"]}')
        print('跳过报告生成和邮件发送，仅启动服务')
        return

    try:
        with open(app.config['LOG_FILE_PATH'], 'r', encoding='utf-8') as f:
            log_content = f.read()

        analyzer = LogAnalyzer(log_content)
        test_cases = analyzer.analyze()

        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        filename = f'engine_log_report_{timestamp}.docx'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)

        generate_word_report(test_cases, output_path)

        recipients = app.config.get('recipients', [])

        if recipients:
            from_email = app.config['mail']['from_email']
            password = app.config['mail']['password']

            if from_email and password:
                local_ip = get_local_ip()
                report_url = f'http://{local_ip}:5000/report/{timestamp}'
                email_result = send_email_report(recipients, output_path, report_url, timestamp)
                print(f'邮件发送结果: {email_result}')
                print(f'报告访问地址: {report_url}')
            else:
                print('邮件配置未设置，跳过邮件发送')
        else:
            print('未配置收件人，跳过邮件发送')

        print(f'报告已生成: {output_path}')
    except Exception as e:
        print(f'启动时生成报告失败: {e}')

if __name__ == '__main__':
    if os.environ.get('WERKZEUG_RUN_MAIN') != 'true':
        print('正在生成报告并发送邮件...')
        generate_startup_report()
        print('报告生成完成，启动 Web 服务...')
    local_ip = get_local_ip()
    print(f'\n访问地址: http://{local_ip}:5000/analysis')
    app.run(host='0.0.0.0', port=5000, debug=True)
